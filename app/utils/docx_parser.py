from pathlib import Path

from docx import Document
from fastapi import HTTPException, status


def extract_text_from_docx(file_path: str | Path) -> str:
    """
    Extract plain text from a DOCX resume.

    Combines paragraph text and table cell text in document order.
    Raises HTTPException (422) for corrupted DOCX files.
    """
    try:
        doc = Document(str(file_path))

        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read the DOCX file. It may be corrupted.",
        )

    return "\n".join(parts)

